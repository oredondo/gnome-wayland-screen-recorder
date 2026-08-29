import os
import re
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger("DocxExporter")


class DocxExporter:
    """Utility class to convert Markdown content or .md files into styled .docx Word documents."""

    @staticmethod
    def markdown_to_docx(md_content: str, output_docx_path: str) -> str:
        """Converts Markdown text into a styled Word (.docx) document."""
        output_dir = os.path.dirname(os.path.abspath(output_docx_path))
        os.makedirs(output_dir, exist_ok=True)

        doc = Document()
        lines = md_content.splitlines()
        in_code_block = False
        code_block_lines = []

        for line in lines:
            stripped = line.strip()

            # Handle Code Blocks
            if stripped.startswith("```"):
                if in_code_block:
                    code_text = "\n".join(code_block_lines)
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.4)
                    run = p.add_run(code_text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(40, 40, 40)
                    in_code_block = False
                    code_block_lines = []
                else:
                    in_code_block = True
                    code_block_lines = []
                continue

            if in_code_block:
                code_block_lines.append(line)
                continue

            # Empty lines
            if not stripped:
                doc.add_paragraph()
                continue

            # Headings
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:].strip(), level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:].strip(), level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:].strip(), level=3)
            elif stripped.startswith("#### "):
                doc.add_heading(stripped[5:].strip(), level=4)
            # Bullet lists
            elif stripped.startswith(("* ", "- ", "+ ")):
                p = doc.add_paragraph(style="List Bullet")
                DocxExporter._add_formatted_text(p, stripped[2:].strip())
            # Numbered lists
            elif re.match(r"^\d+\.\s+", stripped):
                text_part = re.sub(r"^\d+\.\s+", "", stripped)
                p = doc.add_paragraph(style="List Number")
                DocxExporter._add_formatted_text(p, text_part)
            # Blockquotes
            elif stripped.startswith("> "):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                DocxExporter._add_formatted_text(p, stripped[2:].strip())
            # Regular paragraphs
            else:
                p = doc.add_paragraph()
                DocxExporter._add_formatted_text(p, stripped)

        doc.save(output_docx_path)
        logger.info(f"Generated DOCX document: {output_docx_path}")
        return output_docx_path

    @staticmethod
    def convert_file(filepath: str) -> str:
        """Converts an existing .md or .txt file to a .docx file in the same directory."""
        if not os.path.exists(filepath):
            logger.warning(f"DocxExporter: file not found on disk ({filepath}); skipping docx conversion.")
            return ""

        output_docx_path = os.path.splitext(filepath)[0] + ".docx"
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                md_content = f.read()

            return DocxExporter.markdown_to_docx(md_content, output_docx_path)
        except Exception as e:
            logger.warning(f"DocxExporter conversion error for {filepath}: {e}")
            return ""
    
    @staticmethod
    def _add_formatted_text(paragraph, text: str):
        """Parses inline bold (**text**) and italic (*text*) markers into paragraph runs."""
        tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
        for token in tokens:
            if not token:
                continue
            if token.startswith("**") and token.endswith("**"):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("*") and token.endswith("*"):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            elif token.startswith("`") and token.endswith("`"):
                run = paragraph.add_run(token[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(9.5)
            else:
                paragraph.add_run(token)
